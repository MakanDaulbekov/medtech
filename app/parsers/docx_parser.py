"""
DOCX Parser
- Accepts tracked changes: uses accepted (final) text
- Extracts all tables, maps columns heuristically
"""
import re
import logging
from typing import List, Tuple, Optional
from decimal import Decimal, InvalidOperation

import docx
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)


def _clean_price(raw: str) -> Optional[Decimal]:
    if not raw:
        return None
    cleaned = re.sub(r"[^\d.,]", "", str(raw)).replace(",", ".")
    try:
        val = Decimal(cleaned)
        return val if val > 0 else None
    except InvalidOperation:
        return None


def _is_price_col(header: str) -> str | None:
    h = header.lower()
    if "нерезидент" in h or "non" in h:
        return "nonresident"
    if any(k in h for k in ["цена", "стоимость", "резидент", "тариф", "сумма", "kzt", "тг"]):
        return "resident"
    return None


def _is_service_col(header: str) -> bool:
    h = header.lower()
    return any(k in h for k in ["услуга", "наименование", "название", "описание", "процедура"])


def _cell_text(cell) -> str:
    """Get full text including text from tracked-changes runs."""
    parts = []
    for para in cell.paragraphs:
        for run in para.runs:
            parts.append(run.text)
        # Also grab inserted text from tracked changes
        for ins in cell._tc.findall(f".//{qn('w:ins')}"):
            for r in ins.findall(f".//{qn('w:r')}"):
                for t in r.findall(f"{qn('w:t')}"):
                    parts.append(t.text or "")
    return " ".join(parts).strip()


def parse_docx(file_path: str) -> Tuple[List[dict], str]:
    """Returns (items, raw_text)."""
    doc = docx.Document(file_path)
    items: List[dict] = []
    raw_parts: List[str] = []

    # Collect paragraph text
    for para in doc.paragraphs:
        if para.text.strip():
            raw_parts.append(para.text.strip())

    # Process each table
    for table in doc.tables:
        rows_data = []
        for row in table.rows:
            rows_data.append([_cell_text(cell) for cell in row.cells])

        if len(rows_data) < 2:
            continue

        # Find header row (first row with text)
        header_idx = 0
        for i, row in enumerate(rows_data):
            if any(cell.strip() for cell in row):
                header_idx = i
                break

        headers = rows_data[header_idx]

        # Map column indices
        svc_col = None
        code_col = None
        res_col = None
        nonres_col = None

        for j, h in enumerate(headers):
            if _is_service_col(h) and svc_col is None:
                svc_col = j
            if "код" in h.lower() and code_col is None:
                code_col = j
            kind = _is_price_col(h)
            if kind == "resident" and res_col is None:
                res_col = j
            elif kind == "nonresident" and nonres_col is None:
                nonres_col = j

        # If we didn't detect a service col, take col 0 or 1
        if svc_col is None:
            svc_col = 1 if len(headers) > 1 else 0

        for row in rows_data[header_idx + 1:]:
            if not row or all(not c for c in row):
                continue

            name = row[svc_col].strip() if svc_col < len(row) else ""
            if not name or len(name) < 3:
                continue

            code = row[code_col].strip() if code_col is not None and code_col < len(row) else None
            price_r = _clean_price(row[res_col]) if res_col is not None and res_col < len(row) else None
            price_nr = _clean_price(row[nonres_col]) if nonres_col is not None and nonres_col < len(row) else None

            # Last resort: scan all cols for a price
            if price_r is None:
                for j, cell in enumerate(row):
                    if j == svc_col:
                        continue
                    candidate = _clean_price(cell)
                    if candidate and Decimal("1") < candidate < Decimal("10000000"):
                        price_r = candidate
                        break

            items.append({
                "service_name_raw": name,
                "service_code_source": code,
                "price_resident_kzt": price_r,
                "price_nonresident_kzt": price_nr,
                "currency_original": "KZT",
            })

    raw_text = "\n".join(raw_parts)
    return items, raw_text
